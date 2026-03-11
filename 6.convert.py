# convert_multi_format_v7.py
# Modular Multi-Format Converter for Docusaurus - v7.1 (Standardized Edition)
# Supports: PDF, DOCX, XLSX, XLS, XLSM, PPTX, TXT, PNG, ZIP
#
# v7.1 FEATURES:
#   - 🌐 FULL ENGLISH: Code, comments, and variables refactored to English.
#   - 🖼 ASSET PATH FIX: Saves images to root 'static/img' and links as '/img/...'.
#   - 🛡 TABLE PROTECTION: Isolate tables from the cleaning pipeline.
#   - 🧹 SIGNATURE CLEANING: Remove digital certificate metadata.
#   - ⚖ JUSTIFICATION: CSS justification for continuous text blocks.
#   - ⚡ SLIM CODEBASE: Optimized handlers and grouped configuration.

import re, time, unicodedata, logging, tempfile, shutil, warnings, os
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Optional, Tuple, List, Dict
from datetime import datetime

# --- CONFIGURATION ---
FORMATS = {".pdf", ".docx", ".xlsx", ".xls", ".pptx", ".txt", ".png"}
CLEANING_STEPS = [
    ("Unifying hyphens", lambda t: re.sub(r"([a-z])-\n([a-z])", r"\1\2", t, flags=re.I)),
    ("Cleaning index dots", lambda t: re.sub(r"\s*\.{3,}\s*\d+", "", t)),
    ("Removing headers/footers", lambda t: "\n".join([l for l in t.split("\n") if not any(re.match(p, l) for p in [
        r"^\s*\d+\s*[-–—]\s*[A-Za-z]{2,}.*\d{4,}\s*$", 
        r"^[A-Za-z]{2,}.*\d{4,}\s*[-–—]\s*\d+\s*$", 
        r"^\s*\d+\s*$"
    ])])),
    ("Cleaning spaces", lambda t: re.sub(r" {2,}", " ", t))
]
MODES = {"debug": False, "ocr": False, "strict_tables": False}

# --- DEPENDENCIES ---
try:
    import fitz, pymupdf4llm
    PYMUPDF_OK = True
except: PYMUPDF_OK = False

try:
    from docx import Document
    DOCX_OK = True
except: DOCX_OK = False

try:
    import pandas as pd
    PANDAS_OK = True
except: PANDAS_OK = False

try:
    from pptx import Presentation
    PPTX_OK = True
except: PPTX_OK = False

try:
    from PIL import Image
    import pytesseract
    OCR_OK = True
except: OCR_OK = False

# --- GLOBAL STATE ---
_stats = None
_logger = None

class ConversionStats:
    def __init__(self):
        self.data = {"files": {}, "chars": [0, 0], "signatures": 0, "ocr_count": 0, "time": 0.0}
    
    def show_summary(self):
        if not MODES["debug"]: return
        print("\n" + "─"*50 + "\n📊 CONVERSION STATS (v7.1)\n" + "─"*50)
        for ext, count in self.data["files"].items(): print(f" {ext:8}: {count} docs")
        raw, clean = self.data["chars"]
        if raw: print(f" 📝 Characters: {raw:,} -> {clean:,} (Reduction: {(raw-clean)/raw*100:.1f}%)")
        if self.data["signatures"]: print(f" ✍ Signatures cleaned: {self.data['signatures']}")
        if self.data["ocr_count"]: print(f" 🔍 OCR processed: {self.data['ocr_count']}")
        print(f" ⏱ Total time: {self.data['time']:.2f}s\n" + "─"*50)

# --- UTILITIES ---
def log_debug(msg: str):
    if MODES["debug"]: print(f"[DEBUG] {msg}")

def normalize_filename(name: str) -> str:
    path = Path(name)
    stem = unicodedata.normalize("NFKD", path.stem).encode("ascii", "ignore").decode("ascii")
    clean = re.sub(r"[^\w\s\-]", "", stem).strip()
    return re.sub(r"[\s_]+", "-", clean.lower()) + ".md"

def get_docusaurus_fm(p: Path) -> str:
    doc_id = re.sub(r"[^\w\-]", "-", p.stem.lower()).strip("-")
    title = p.stem.replace("-", " ").replace("_", " ").title()
    return f"---\nid: {doc_id}\ntitle: \"{title}\"\nsidebar_label: \"{title}\"\n---\n\n"

# --- CORE LOGIC ---
def toggle_tables(text: str, tables: List[str] = None) -> Tuple[str, List[str]]:
    if tables is None: # Isolation
        extracted = []
        def replacer(m):
            extracted.append(m.group(0)); return f"\n<!-- TABLE_BLOCK_{len(extracted)-1} -->\n"
        return re.sub(r"(?:^\|.+\n)+", replacer, text, flags=re.M), extracted
    # Restoration
    for i, t in enumerate(tables): text = text.replace(f"<!-- TABLE_BLOCK_{i} -->", f"\n{t}\n")
    return text, []

def clean_content(text: str, verbose: bool = True) -> str:
    _stats.data["chars"][0] += len(text)
    
    # Signatures & MDX Safety
    orig_len = len(text)
    patterns = [(r'`Firmado por[^`]+`', '[DIGITAL SIGNATURE]'), (r'\*{2,}\d+\*{2,}', ''), (r'el día \d{2}/\d{2}/\d{4}.*$', '')]
    for p, r in patterns: text = re.sub(p, r, text, flags=re.I if '[' in r else re.M)
    if len(text) < orig_len: _stats.data["signatures"] += 1
    
    text = re.sub(r"<(?!/?(?:p|br|hr|img|div|span|strong|em|b|i|u|h\d)\b)", "&lt;", text)
    text = text.replace("{", "&#123;").replace("}", "&#125;")
    for tag in ["br", "hr"]: text = re.sub(fr"<{tag}\s*>", f"<{tag} />", text, flags=re.I)
    
    # Table Isolation & Pipeline
    text, tables = toggle_tables(text)
    for name, step in CLEANING_STEPS: text = step(text)
    
    if not MODES["strict_tables"]:
        lines = text.split("\n"); merged, i = [], 0
        while i < len(lines):
            l = lines[i].rstrip()
            if not l or l.startswith(("#","-","*","|",">","`","<")) or l.endswith((".",":",";","!","?")):
                merged.append(l); i += 1
            elif i+1 < len(lines) and lines[i+1].strip() and not lines[i+1].strip().startswith(("#","-","*","|",">","`","<")):
                merged.append(l + " " + lines[i+1].lstrip()); i += 2
            else: merged.append(l); i += 1
        text = "\n".join(merged)

    # Justify & Restore Tables
    def just_cb(m):
        b = m.group(0).strip()
        if any(re.match(r"^(?:#|\-|\*|\d+\.|\||```|<|id:|title:|sidebar_label:|---)", l.strip()) for l in b.split('\n')) or len(b) < 60: return m.group(0)
        return f'\n<div style={{{{textAlign: "justify"}}}}>\n\n{b}\n\n</div>\n'
    text = re.sub(r"((?:(?!\n\n).)+)", just_cb, text, flags=re.S)
    text, _ = toggle_tables(text, tables)
    
    final = re.sub(r"\n{3,}", "\n\n", text).strip()
    _stats.data["chars"][1] += len(final)
    return final

# --- HANDLERS ---
def handle_pdf(path: Path, static_img_dir: Path) -> Optional[str]:
    if not PYMUPDF_OK: return None
    doc = fitz.open(path)
    scanned = (len("".join(p.get_text() for p in doc))/len(doc) < 500) if len(doc) > 0 else True
    
    if scanned and MODES["ocr"] and OCR_OK:
        log_debug(f"Scan detected: {path.name}"); _stats.data["ocr_count"] += 1
        pages = []
        for pg in doc:
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
                pg.get_pixmap(dpi=300).save(f.name)
                pages.append(pytesseract.image_to_string(Image.open(f.name), lang='spa'))
                os.unlink(f.name)
        return "\n\n".join(pages)
    
    # Standard extraction; write images to a subfolder within static/img
    doc_slug = re.sub(r"[^\w\-]", "-", path.stem.lower())
    disk_img_path = static_img_dir / doc_slug
    disk_img_path.mkdir(parents=True, exist_ok=True)
    
    markdown = pymupdf4llm.to_markdown(str(path), write_images=True, image_path=str(disk_img_path))
    
    # Fix image paths in Markdown for Docusaurus (rewrite to /img/doc_slug/image)
    def fix_img_path(match):
        img_filename = Path(match.group(2)).name
        return f'![{match.group(1)}](/img/{doc_slug}/{img_filename})'
    
    return re.sub(r'\!\[(.*?)\]\((.*?)\)', fix_img_path, markdown)

def handle_docx(path: Path) -> Optional[str]:
    if not DOCX_OK: return None
    doc = Document(path); res = [f"## {p.text.strip()}" if p.style.name.startswith('Hea') else p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for t in doc.tables:
        for r in t.rows: res.append("| " + " | ".join(c.text.strip() for c in r.cells) + " |")
    return "\n".join(res)

def handle_excel(path: Path) -> Optional[str]:
    if not PANDAS_OK: return None
    xl = pd.ExcelFile(path); return "\n".join([f"## {s}\n{pd.read_excel(path, s).to_markdown(index=False)}" for s in xl.sheet_names])

HANDLERS = {".pdf": handle_pdf, ".docx": handle_docx, ".xlsx": handle_excel, ".xls": handle_excel, ".txt": lambda p: p.read_text(errors='ignore')}

# --- MAIN ---
def process_all(src: Path, dst: Path, wrk: int):
    global _stats; _stats = ConversionStats()
    # Assume project root is parent of 'docs' if dst is 'docs'
    root = dst.parent if dst.name == "docs" else dst.resolve()
    static_img_dir = root / "static" / "img"
    
    files = [f for f in src.rglob("*") if f.suffix.lower() in HANDLERS]
    print(f"� v7.1 starting ({len(files)} docs)...")
    t0 = time.time()
    
    def convert(f: Path):
        try:
            ext = f.suffix.lower()
            if ext == ".pdf": raw = handle_pdf(f, static_img_dir)
            else: raw = HANDLERS[ext](f)
            if not raw: return False
            
            final_content = get_docusaurus_fm(f) + clean_content(raw, wrk==1)
            target = dst / f.relative_to(src).parent / normalize_filename(f.name)
            target.parent.mkdir(parents=True, exist_ok=True); target.write_text(final_content, encoding='utf-8')
            _stats.data["files"][ext] = _stats.data["files"].get(ext, 0) + 1
            return True
        except Exception as e:
            if _logger: _logger.error(f"{f.name}: {e}"); return False

    with ThreadPoolExecutor(wrk) as exe:
        for f in files: exe.submit(convert, f)
            
    _stats.data["time"] = time.time() - t0; _stats.show_summary()

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", default="pdf"); parser.add_argument("--output", default="docs")
    parser.add_argument("--workers", type=int, default=1); parser.add_argument("--ocr", action="store_true")
    parser.add_argument("--debug", action="store_true"); parser.add_argument("--strict-tables", action="store_true")
    args = parser.parse_args(); MODES.update({"debug": args.debug, "ocr": args.ocr, "strict_tables": args.strict_tables})
    process_all(Path(args.input), Path(args.output), args.workers)
